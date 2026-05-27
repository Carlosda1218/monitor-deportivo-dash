"""
Genera TECHNICAL_OVERVIEW_CombatIQ.docx con estilo profesional azul (tipo HENKO).
Ejecutar: python _gen_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ────────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1A, 0x3F, 0xA8)   # azul principal
BLUE_LIGHT = RGBColor(0xF0, 0xF3, 0xFB)   # fondo filas alternas
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK       = RGBColor(0x2C, 0x2C, 0x2C)
GRAY       = RGBColor(0x88, 0x88, 0x88)
MONO_BG    = RGBColor(0xEE, 0xF1, 0xFB)   # fondo diagrama
MONO_TEXT  = RGBColor(0x1A, 0x3F, 0xA8)


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Rellena el fondo de una celda con color hexadecimal (sin #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_hr(doc):
    """Añade una línea separadora horizontal."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D8DCE8")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


def h1(doc, text):
    """Título principal (portada)."""
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(36)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after  = Pt(8)


def subtitle(doc, text):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold      = False
    run.font.size = Pt(16)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(30)


def meta(doc, text):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(50)


def h2(doc, text):
    """Encabezado de sección (azul, mayúsculas)."""
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(8)


def h3(doc, text):
    """Encabezado de subsección (azul, normal)."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def h4(doc, text):
    """Encabezado de nivel 4 (oscuro)."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)


def body(doc, text):
    p   = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)


def numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)


def mono(doc, text):
    """Bloque de diagrama/código con fondo azul claro."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = MONO_TEXT
    # fondo del párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F3FB")
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.4)


def add_table(doc, headers, rows):
    """Tabla con cabecera azul y filas alternas."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabecera
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_bg(cell, "1A3FA8")

    # Filas
    for ri, row in enumerate(rows):
        bg = "F0F3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            set_cell_bg(cell, bg)

    doc.add_paragraph()   # espacio tras tabla


# ── Construcción del documento ─────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Fuente base del documento
doc.styles["Normal"].font.name = "Segoe UI"
doc.styles["Normal"].font.size = Pt(10.5)
doc.styles["Normal"].font.color.rgb = DARK


# ══════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════
h1(doc, "CombatIQ")
subtitle(doc, "Technical Overview")
meta(doc, "Versión 1.0  ·  Mayo 2026\nPlataforma de Monitoreo para Deportes de Combate")
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════
h2(doc, "1. System Overview")
body(doc,
    "CombatIQ es una plataforma web de monitoreo del rendimiento diseñada para "
    "deportistas y coaches de artes marciales de combate (Taekwondo y Boxeo). "
    "Integra datos biomédicos de sensores, check-ins de bienestar subjetivo e "
    "inteligencia artificial en un único sistema multi-rol accesible desde "
    "cualquier navegador.")

h3(doc, "Funcionalidades Principales")
bullet(doc, "Monitoreo de señales biomédicas — análisis de ECG/HRV (frecuencia cardíaca y variabilidad) e IMU (detección de golpes, patadas e impactos recibidos).")
bullet(doc, "Check-in de bienestar inteligente — cuestionarios ponderados por deporte que calculan un Wellness Score (0–100) pre-entrenamiento.")
bullet(doc, "Dashboard del coach — vista de equipo en tiempo real con alertas automáticas, historial de sesiones y métricas de carga.")
bullet(doc, "Análisis de carga interna y externa — ACWR (ratio carga aguda:crónica), readiness HRV, tendencias de volumen IMU y alertas por niveles (ok / warning / danger).")
bullet(doc, "Insights narrativos con IA — integración con la API de Claude (Anthropic) que genera análisis de coaching contextualizados al deporte y a los datos del atleta.")
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 2. ARQUITECTURA
# ══════════════════════════════════════════════════════════════════════
h2(doc, "2. Arquitectura del Sistema")
mono(doc,
"┌─────────────────────────────────────────────────────────┐\n"
"│                      USUARIO                            │\n"
"│     Navegador Web (Chrome / Safari / Firefox)           │\n"
"│             Compatible con PWA (instalable)             │\n"
"└──────────────────────┬──────────────────────────────────┘\n"
"                       │  HTTP / HTTPS\n"
"┌──────────────────────▼──────────────────────────────────┐\n"
"│                     FRONTEND                            │\n"
"│  Plotly Dash (Python) + CSS personalizado + JavaScript  │\n"
"│  Layouts reactivos · Gráficas interactivas · DataTable  │\n"
"└──────────────────────┬──────────────────────────────────┘\n"
"                       │  Callbacks (Dash) + Rutas (Flask)\n"
"┌──────────────────────▼──────────────────────────────────┐\n"
"│                     BACKEND                             │\n"
"│           Flask (Python) + Dash Server                  │\n"
"│  Autenticación · Sesiones · API REST · Archivos         │\n"
"│  analysis_engine.py · ai_insights.py (Claude API)       │\n"
"└──────────────────────┬──────────────────────────────────┘\n"
"                       │  SQL (sqlite3)\n"
"┌──────────────────────▼──────────────────────────────────┐\n"
"│                  BASE DE DATOS                          │\n"
"│        SQLite — data/users.db — modo WAL                │\n"
"│  20+ tablas · Migraciones versionadas (v10 → v210)      │\n"
"└─────────────────────────────────────────────────────────┘")

add_table(doc,
    ["Capa", "Tecnología", "Responsabilidad"],
    [
        ["Usuario",      "Navegador web / PWA",                      "Interfaz visual, instalable en móvil como Progressive Web App"],
        ["Frontend",     "Plotly Dash 2.16, Plotly 6.4, CSS, JS",    "Componentes reactivos, gráficas biomédicas, DataTables, sliders de check-in"],
        ["Backend",      "Flask 3.0, Python 3.11+",                  "Servidor WSGI, autenticación bcrypt, rutas REST API, carga de archivos CSV/JSON"],
        ["Base de datos","SQLite (WAL mode)",                         "Persistencia de usuarios, métricas, cuestionarios, sesiones; migraciones automáticas"],
        ["IA externa",   "Anthropic Claude API (anthropic==0.96)",    "Generación de notas narrativas de coaching basadas en el informe del atleta"],
    ]
)
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 3. COMPONENTES
# ══════════════════════════════════════════════════════════════════════
h2(doc, "3. Componentes del Sistema")

h4(doc, "app.py")
body(doc,
    "Punto de entrada principal. Inicializa el servidor Flask y la instancia Dash, "
    "registra el layout base multi-rol (deportista / coach) y define todos los callbacks "
    "reactivos. Gestiona la autenticación, la navegación entre páginas y orquesta las "
    "cuatro vistas modulares: SignalsView, SensorsView, CompareView y AnalysisView. "
    "También expone los endpoints REST internos (/api/sensor-data, /api/sensor-ping) "
    "para ingestión de datos desde dispositivos externos.")

h4(doc, "db.py")
body(doc,
    "Capa completa de acceso a datos. Define todas las tablas SQLite y ejecuta migraciones "
    "versionadas de la v10 a la v210 de forma automática e idempotente al arrancar la app. "
    "Expone funciones CRUD para: usuarios, sesiones de entrenamiento, cuestionarios, "
    "métricas de sensores (ECG, IMU, EMG, RESP), equipos, mensajes internos, peso, "
    "nutrición y recuperación de contraseña mediante tokens temporales con expiración.")

h4(doc, "sensors.py")
body(doc,
    "Catálogo declarativo de los sensores del sistema. Cada entrada define nombre, "
    "señales que produce, métricas calculadas, compatibilidad por deporte y ejemplos "
    "de hardware compatible. Los sensores disponibles son:")
add_table(doc,
    ["Código", "Sensor", "Deporte", "Métricas clave"],
    [
        ["ECG",       "Banda torácica ECG / HRV", "TKD · Boxeo", "BPM, SDNN, RMSSD, n_peaks"],
        ["IMU_WRIST", "IMU muñeca / guante",      "TKD · Boxeo", "golpes/min, intensidad media (g), intensidad máxima (g)"],
        ["IMU_FOOT",  "IMU tobillo / pie",         "TKD",         "patadas/min, potencia (g), ratio dominante/no-dominante"],
        ["IMU_HEAD",  "IMU casco / cabeza",        "TKD · Boxeo", "impactos recibidos, pico de g, impactos >3 g, impactos >6 g"],
        ["HR_WRIST",  "FC reloj / pulsera",        "TKD · Boxeo", "BPM promedio y máximo"],
    ]
)

h4(doc, "questionnaires.py")
body(doc,
    "Define 22 preguntas organizadas en 7 grupos temáticos: base, taekwondo, boxeo, "
    "competencia, peso y molestia. Cada pregunta tiene un peso numérico (3–14) y una "
    "dimensión (positive o risk). El módulo implementa el algoritmo de cálculo del "
    "Wellness Score (0–100): normaliza cada respuesta al rango 0–100, invierte las "
    "dimensiones de riesgo y pondera cada ítem según su peso relativo al deporte del atleta.")

h4(doc, "analysis_engine.py")
body(doc,
    "Motor de análisis que agrega datos históricos del atleta y genera un informe "
    "estructurado. Calcula: ACWR (ratio carga aguda:crónica de las últimas 7 vs 28 "
    "sesiones), readiness HRV comparada contra la baseline rolling de 30 días, "
    "tendencias de volumen IMU y alertas automáticas clasificadas por nivel "
    "(ok, warning, danger). Incluye caché en memoria con TTL de 5 minutos.")

h4(doc, "ai_insights.py")
body(doc,
    "Integración con la API de Claude de Anthropic. Recibe el informe del "
    "analysis_engine y genera un análisis narrativo de coaching adaptado al deporte "
    "y al momento de la temporada. Implementa caché en memoria (10 min), timeout "
    "configurable (defecto 8 s) y degradación elegante: si no hay API key o falla "
    "la llamada, retorna un mensaje de aviso sin romper la UI.")
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 4. FLUJO DE DATOS
# ══════════════════════════════════════════════════════════════════════
h2(doc, "4. Flujo de Datos")

h3(doc, "Flujo 1: Check-in de bienestar (deportista)")
numbered(doc, "El deportista autenticado navega a la página /wellbeing.")
numbered(doc, "El layout detecta su deporte (users.sport) y muestra solo las preguntas relevantes (preguntas base + grupo específico del deporte).")
numbered(doc, "El deportista responde los sliders de 1–5 (energía, recuperación, sueño, fatiga, etc.) y registra RPE y duración de la sesión anterior.")
numbered(doc, "questionnaires.compute_wellness_score(answers, sport) normaliza cada respuesta a 0–100, invierte las dimensiones de riesgo y aplica los pesos configurados.")
numbered(doc, "El score final (0–100) y el JSON de respuestas se persisten en la tabla questionnaires junto al user_id, timestamp y session_id.")
numbered(doc, "analysis_engine.invalidate_cache(uid) limpia el caché del atleta.")
numbered(doc, "El dashboard del coach actualiza la vista del equipo: el wellness del atleta aparece con su nivel de color (verde / naranja / rojo).")
numbered(doc, "Si ANTHROPIC_API_KEY está configurada, ai_insights.generate_coaching_note() devuelve el análisis narrativo visible para el coach.")

h3(doc, "Flujo 2: Análisis de señal ECG")
numbered(doc, "El coach o deportista sube el CSV de la banda torácica en la vista Signals.")
numbered(doc, "El backend lee el archivo, detecta la frecuencia de muestreo y registra el archivo en ecg_files.")
numbered(doc, "analysis_engine carga las muestras y aplica scipy.signal.find_peaks() para detectar los picos R de la señal.")
numbered(doc, "Se calculan los intervalos R–R y a partir de ellos: BPM, SDNN y RMSSD.")
numbered(doc, "Las métricas se guardan en ecg_metrics vinculadas al ecg_file_id y session_id.")
numbered(doc, "La vista Signals renderiza la señal ECG con los picos marcados y las métricas en tarjetas de resumen.")
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 5. MODELO DE DATOS
# ══════════════════════════════════════════════════════════════════════
h2(doc, "5. Modelo de Datos")
h3(doc, "Tablas principales")
add_table(doc,
    ["Tabla", "Columnas clave", "Descripción"],
    [
        ["users",
         "id, name, email, role, sport, password_hash, coach_id, avatar_url, onboarding_done",
         "Usuarios del sistema. Roles: deportista / coach. coach_id enlaza atleta con su coach."],
        ["questionnaires",
         "id, user_id, ts, answers_json, wellness_score, rpe, duration_min, session_id",
         "Check-ins diarios. answers_json guarda las respuestas; wellness_score es el valor ponderado 0–100."],
        ["sessions",
         "id, athlete_id, created_by, ts_start, ts_end, sport, notes, status",
         "Sesiones de entrenamiento. Agrupa ECG, IMU y cuestionario bajo una misma entidad temporal."],
        ["ecg_files",
         "id, user_id, filename, fs, created_at, session_id",
         "Archivos CSV de señal ECG. fs = frecuencia de muestreo en Hz."],
        ["ecg_metrics",
         "id, ecg_file_id, bpm, sdnn, rmssd, n_peaks, created_at",
         "Métricas calculadas sobre cada archivo ECG. Relación N:1 con ecg_files."],
        ["imu_metrics",
         "id, user_id, filename, ts, n_hits, hits_per_min, mean_int_g, max_int_g, sensor_type, mean_ang_vel, max_ang_vel, session_id",
         "Métricas de impacto de cada sensor IMU. sensor_type indica muñeca, pie o cabeza."],
        ["coach_athletes",
         "coach_id, athlete_id, created_at",
         "Relación de adopción coach–atleta (modelo nuevo). Complementa el coach_id legacy."],
        ["teams / team_members",
         "team_id, coach_id, name, sport / team_id, athlete_id, role_label",
         "Gestión de equipos. Un coach puede tener múltiples equipos con atletas asignados."],
        ["weights",
         "id, user_id, date, weight_kg, target_kg",
         "Historial de peso con peso objetivo para control de categoría de competencia."],
        ["messages",
         "id, sender_id, receiver_id, body, ts, read_at",
         "Chat interno coach–atleta. read_at indica si el mensaje fue leído."],
    ]
)

h3(doc, "Relación entre tablas principales")
body(doc,
    "Un usuario (role='deportista') pertenece a un coach vía coach_athletes (relación de "
    "adopción) o por el campo legacy users.coach_id. Cada session agrupa las métricas del "
    "entrenamiento: enlaza con ecg_files, imu_metrics y questionnaires mediante session_id.")
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 6. PROCESAMIENTO BIOMÉDICO
# ══════════════════════════════════════════════════════════════════════
h2(doc, "6. Procesamiento Biomédico")

h3(doc, "Señal: ECG (electrocardiograma)")
add_table(doc,
    ["Parámetro", "Significado", "Cálculo"],
    [
        ["BPM",    "Frecuencia cardíaca (latidos por minuto)",
         "Se detectan picos R con scipy.signal.find_peaks() sobre la señal filtrada. BPM = 60 / media(intervalos R–R en segundos)."],
        ["SDNN",   "Variabilidad HRV global",
         "Desviación estándar de todos los intervalos R–R de la grabación. Valores bajos indican fatiga acumulada."],
        ["RMSSD",  "HRV de corto plazo (actividad parasimpática)",
         "Raíz del promedio de los cuadrados de las diferencias entre intervalos R–R consecutivos. Indicador de recuperación autonómica."],
        ["n_peaks","Total de latidos detectados",
         "Contador de picos R identificados en el segmento analizado. Sirve para validar la calidad de la señal."],
    ]
)

h3(doc, "Señal: IMU (unidad de medición inercial)")
add_table(doc,
    ["Parámetro", "Significado", "Cálculo"],
    [
        ["n_hits",
         "Total de golpes / patadas / impactos",
         "Magnitud vectorial: √(ax²+ay²+az²). Los picos ≥ 2 g se cuentan como impactos."],
        ["hits_per_min",
         "Cadencia (golpes por minuto)",
         "n_hits / duración en minutos. Indicador de volumen de trabajo específico."],
        ["mean_int_g / max_int_g",
         "Intensidad media y máxima del impacto",
         "Valor de pico de la magnitud en cada impacto detectado, expresado en g (gravedades)."],
        ["mean_ang_vel / max_ang_vel",
         "Velocidad angular media y máxima",
         "Derivada del giroscopio triaxial. Cuantifica la rapidez de rotación del segmento corporal durante el impacto."],
    ]
)
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 7. LIMITACIONES
# ══════════════════════════════════════════════════════════════════════
h2(doc, "7. Limitaciones")
numbered(doc, "Base de datos SQLite: concurrencia limitada bajo cargas altas simultáneas (múltiples coaches + atletas en la misma instancia). El modo WAL mejora la situación pero no la elimina.")
numbered(doc, "Sin integración BLE en tiempo real: la ingestión de ECG e IMU es mayoritariamente manual (CSV/JSON upload). El hub BLE Nordic está en fase de prototipo.")
numbered(doc, "Sin app móvil nativa: la plataforma es solo web (PWA instalable desde el navegador). No existe versión iOS o Android publicada en tiendas de apps.")
numbered(doc, "Insights de IA opcionales: requieren ANTHROPIC_API_KEY activa. Sin ella la sección de análisis narrativo no se muestra, pero el resto del sistema funciona sin modificaciones.")
numbered(doc, "Análisis biomecánico por video (YOLO / MediaPipe) computacionalmente intenso: no es apto para servidores sin GPU. Las estimaciones de pose se realizan offline sobre archivos de video previamente grabados.")
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 8. MEJORAS
# ══════════════════════════════════════════════════════════════════════
h2(doc, "8. Mejoras")

h3(doc, "Arquitectura")
body(doc,
    "Migrar a una arquitectura API REST + SPA (FastAPI en backend, React o Vue en frontend) "
    "para desacoplar las capas y escalar cada una de forma independiente. Introducir Redis "
    "como caché compartida para los informes del analysis_engine, eliminando el caché en "
    "memoria no compartido entre procesos.")

h3(doc, "Base de datos")
body(doc,
    "Migrar de SQLite a PostgreSQL cuando el número de atletas activos supere los 300 o "
    "cuando se despliegue en producción con múltiples coaches concurrentes. Adoptar "
    "SQLAlchemy como ORM para simplificar las migraciones (Alembic) y mejorar la cobertura "
    "de tests. La capa de migraciones versionadas que ya existe en db.py puede reutilizarse "
    "para la transición sin pérdida de datos.")

h3(doc, "Funcionalidad")
body(doc,
    "Integración BLE en tiempo real con transmisión desde wearables directamente a la "
    "plataforma sin pasos manuales de carga. Notificaciones push al móvil del coach cuando "
    "un atleta registra un wellness score bajo o un impacto de alto riesgo. App móvil nativa "
    "(React Native) para el check-in diario del deportista sin necesidad de abrir el navegador.")
add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 9. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════
h2(doc, "9. Conclusiones")

h3(doc, "Qué has aprendido")
bullet(doc, "Cómo construir un sistema web multi-rol real (coach / deportista) con autenticación segura (bcrypt + tokens de sesión), navegación protegida y separación de permisos por rol.")
bullet(doc, "Procesamiento de señales biomédicas en Python: detección de picos R en ECG con scipy, cálculo de parámetros HRV (SDNN, RMSSD) y detección de impactos en señales IMU triaxiales.")
bullet(doc, "Diseño de un esquema de base de datos evolutivo con migraciones versionadas que permiten añadir funcionalidades en producción sin perder datos existentes.")
bullet(doc, "Integración de modelos de IA generativa (Claude API) en el flujo de datos de una aplicación real, con manejo de errores, caché y degradación elegante cuando el servicio no está disponible.")
bullet(doc, "Patrones de arquitectura reactiva con Plotly Dash: callbacks complejos con múltiples inputs/outputs, estados compartidos y componentes dinámicos por deporte.")

h3(doc, "Qué mejorarías")
bullet(doc, "Separar el frontend del backend desde el inicio (API REST + cliente independiente) para facilitar el testing y la escala horizontal.")
bullet(doc, "Implementar WebSockets o Server-Sent Events para que el dashboard del coach se actualice en tiempo real sin necesidad de recargar la página.")
bullet(doc, "Añadir una suite de tests automatizados (pytest + coverage) desde el primer día, especialmente para el Wellness Score y el pipeline de análisis ECG/IMU.")
bullet(doc, "Usar PostgreSQL desde el despliegue inicial en lugar de migrar después, para evitar trabajo de conversión de datos en producción.")

add_hr(doc)


# ══════════════════════════════════════════════════════════════════════
# 10. FORTALEZAS COMPETITIVAS — EN PALABRAS SIMPLES
# ══════════════════════════════════════════════════════════════════════
h2(doc, "10. Fortalezas Competitivas — En Palabras Simples")
body(doc,
    "Este apartado explica por qué CombatIQ es diferente a lo que existe hoy, "
    "sin tecnicismos. Está pensado para que cualquier persona, sin importar su "
    "nivel técnico, entienda el valor de lo que se ha construido.")

h3(doc, "¿Qué problema resuelve?")
body(doc, "Un coach de Taekwondo o Boxeo trabaja hoy así:")
bullet(doc, "Le pregunta al atleta \"¿cómo te sientes?\" y anota la respuesta en papel o en un chat de WhatsApp.")
bullet(doc, "Sube datos de un sensor a una computadora, abre Excel y trata de interpretar números sin contexto.")
bullet(doc, "No tiene forma de comparar cómo estaba el atleta hace tres semanas con cómo está hoy.")
bullet(doc, "Si dirige un equipo de 15 personas, esa información está dispersa en 15 conversaciones distintas.")
body(doc, "CombatIQ reemplaza todo eso con una sola plataforma, ordenada y automática.")

h3(doc, "Lo que hace CombatIQ — sin tecnicismos")
body(doc,
    "Antes de cada entrenamiento, el deportista responde en su celular unas preguntas "
    "breves: ¿cómo dormiste?, ¿qué tan cansado llegas?, ¿sientes alguna molestia? "
    "La app convierte esas respuestas en un número del 0 al 100 —el Wellness Score— "
    "que el coach ve al instante en su panel.")
body(doc,
    "Si el atleta también usó una banda en el pecho durante el entrenamiento, la app "
    "lee la señal del corazón y calcula automáticamente la frecuencia cardíaca y la "
    "variabilidad (HRV), que son los indicadores más precisos para saber si un "
    "deportista está recuperado o sobreentrenado.")
body(doc,
    "Si además llevó un sensor en la muñeca o el tobillo, la app cuenta cuántos golpes "
    "dio, a qué ritmo y con qué fuerza. Sin que nadie tenga que contar nada manualmente.")
body(doc,
    "Y al final, la inteligencia artificial —la misma tecnología que hay detrás de "
    "ChatGPT— lee todos esos datos juntos y le escribe al coach un párrafo en lenguaje "
    "natural explicando el estado del atleta y qué cuidados tiene que tener en el "
    "siguiente entrenamiento.")

h3(doc, "Por qué es difícil de copiar")
body(doc,
    "No es solo una app de cuestionarios, ni solo un visualizador de sensores, ni solo "
    "un chat con IA. CombatIQ combina las tres cosas en un sistema coherente donde cada "
    "dato se entiende en contexto de los demás.")
body(doc,
    "Además, está construida específicamente para deportes de combate. Las preguntas, "
    "los umbrales de alerta, las métricas calculadas y el lenguaje que usa la IA están "
    "pensados para Taekwondo y Boxeo —no para atletismo ni para fútbol. Eso hace que "
    "el coach sienta que la herramienta entiende su deporte.")

h3(doc, "Comparación con lo que existe hoy")
add_table(doc,
    ["¿Qué tienen hoy?", "¿Qué problema tiene?", "¿Cómo lo resuelve CombatIQ?"],
    [
        ["Papel / WhatsApp",
         "Sin historial, sin análisis, la información se pierde.",
         "Registro digital automático con historial completo por sesión y por atleta."],
        ["Polar / Garmin solos",
         "Solo miden ritmo cardíaco, sin contexto de entrenamiento ni rol de coach.",
         "CombatIQ añade contexto: bienestar subjetivo + impactos + análisis de IA encima del dato de la banda."],
        ["TrainingPeaks / Final Surge",
         "No son para deportes de combate; sin IMU ni detección de golpes.",
         "Cuestionarios, alertas y métricas calibradas específicamente para TKD y Boxeo."],
        ["Excel del profe",
         "Manual, tardado, sin alertas automáticas, sin tendencias.",
         "Automatización completa: los datos entran solos, las alertas salen solas."],
    ]
)

h3(doc, "Lo que hace que esto valga dinero")
numbered(doc,
    "Los datos son del equipo. A diferencia de apps de terceros donde la información "
    "queda en servidores ajenos, aquí el equipo es dueño de su propio historial de rendimiento.")
numbered(doc,
    "Prevención de lesiones. El sistema detecta patrones de sobrecarga antes de que el "
    "atleta se lastime, no después. Un atleta lesionado es tiempo perdido y resultados perdidos.")
numbered(doc,
    "Decisiones objetivas. El coach deja de adivinar si el atleta está listo para una "
    "sesión exigente. Tiene un número, una tendencia y un análisis escrito frente a él cada mañana.")
numbered(doc,
    "Escala sin esfuerzo extra. Da lo mismo si el equipo tiene 5 atletas o 50. El coach "
    "ve el estado de todos en una sola pantalla, sin abrir 50 chats.")
numbered(doc,
    "Tecnología de vanguardia accesible. El uso de inteligencia artificial para generar "
    "análisis de coaching no existe en ninguna herramienta específica para deportes de "
    "combate en el mercado hispanohablante.")

add_hr(doc)

# Pie de página
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documento elaborado en mayo de 2026  ·  CombatIQ — Plataforma de Monitoreo para Deportes de Combate")
run.font.size = Pt(8.5)
run.font.color.rgb = GRAY


# ── Guardar ────────────────────────────────────────────────────────────────────
import os, shutil
out_tmp  = "TECHNICAL_OVERVIEW_CombatIQ_new.docx"
out_final = "TECHNICAL_OVERVIEW_CombatIQ.docx"
doc.save(out_tmp)
# Reemplaza el final solo si el tmp se generó bien
if os.path.exists(out_final):
    try:
        os.remove(out_final)
    except PermissionError:
        print("AVISO: cierra el Word antes de reemplazar el archivo original.")
        print("Archivo guardado como: " + out_tmp)
        raise SystemExit(0)
shutil.move(out_tmp, out_final)
print("Archivo generado: " + out_final)
