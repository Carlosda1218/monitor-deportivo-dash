"""
Genera 'Alcance_CombatIQ.docx' con el mismo formato profesional que el
documento Plan_Proyecto_Henko.pdf (portada, secciones numeradas, tablas azules,
encabezado y pie con paginacion).

Contenido: vision general de producto para inversor / direccion. Explica QUE es
la app, su alcance funcional, el analisis que hay detras, el hardware y la
escalabilidad. No incluye propuesta economica.

Ejecutar:  python _gen_alcance_combatiq.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colores (paleta Henko) ──────────────────────────────────────────────────
BLUE       = RGBColor(0x1A, 0x3F, 0xA8)   # azul principal
RED        = RGBColor(0xB0, 0x2A, 0x2A)   # cabecera de riesgos
GREEN      = RGBColor(0x2E, 0x7D, 0x32)   # nivel ok
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK       = RGBColor(0x2C, 0x2C, 0x2C)
GRAY       = RGBColor(0x88, 0x88, 0x88)
MONO_TEXT  = RGBColor(0x1A, 0x3F, 0xA8)


# ── Helpers de estilo ───────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_hr(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D8DCE8")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)


def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(40)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after  = Pt(8)


def subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(17)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)


def cover_italic(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(40)


def meta(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)


def meta_soft(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(2)


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(8)


def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)


def mono(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = MONO_TEXT
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F3FB")
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.4)


def add_table(doc, headers, rows, header_hex="1A3FA8", col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        set_cell_bg(cell, header_hex)
    for ri, row in enumerate(rows):
        bg = "F0F3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK
            set_cell_bg(cell, bg)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for r in t.rows:
                r.cells[ci].width = Cm(w)
    doc.add_paragraph()


def section_break(doc):
    """Salto de pagina."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ── Encabezado + pie con paginacion ─────────────────────────────────────────

def _field(paragraph, instr):
    run = paragraph.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
    it  = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    run._r.append(fc1); run._r.append(it); run._r.append(fc2)
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY


def setup_header_footer(doc):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True   # portada sin encabezado/pie

    # Encabezado (paginas != portada)
    hdr = sec.header
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("CombatIQ · Vision de producto")
    hr.font.size = Pt(8.5)
    hr.font.color.rgb = GRAY

    # Pie con "Pagina X de Y"
    ftr = sec.footer
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = fp.add_run("Pagina "); r0.font.size = Pt(8.5); r0.font.color.rgb = GRAY
    _field(fp, "PAGE")
    r1 = fp.add_run(" de "); r1.font.size = Pt(8.5); r1.font.color.rgb = GRAY
    _field(fp, "NUMPAGES")


# ══════════════════════════════════════════════════════════════════════════
# CONSTRUCCION DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.6)
    section.right_margin  = Cm(2.6)

doc.styles["Normal"].font.name = "Segoe UI"
doc.styles["Normal"].font.size = Pt(10.5)
doc.styles["Normal"].font.color.rgb = DARK

setup_header_footer(doc)

# ── PORTADA ──────────────────────────────────────────────────────────────────
h1(doc, "CombatIQ")
subtitle(doc, "Plataforma digital de monitoreo para deportes de combate")
cover_italic(doc, "Alcance funcional · Cómo funciona el análisis · Hardware · Escalabilidad")
meta(doc, "Documento de visión general de producto")
meta(doc, "Taekwondo y Boxeo · Atleta · Coach · Dirección")
meta_soft(doc, "Versión 1.0 · Mayo 2026 · En producción")
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "1. Resumen ejecutivo")
body(doc,
    "CombatIQ es una plataforma web que reúne, en un solo lugar, todo lo que un "
    "atleta y un coach de deportes de combate necesitan para entrenar con datos y "
    "no a ciegas: el bienestar diario del deportista, las señales de sus sensores "
    "(corazón e impactos), el análisis de su técnica en video y un asistente de "
    "inteligencia artificial que traduce todo eso en recomendaciones concretas.")
body(doc,
    "Está pensada específicamente para Taekwondo y Boxeo —no es una app de fitness "
    "genérica—. Las preguntas, las métricas, los umbrales de alerta y el lenguaje "
    "que usa la IA están calibrados para esos deportes. Hoy el foco está en "
    "Taekwondo, donde la plataforma fue validada con video y atletas reales; la "
    "arquitectura ya reconoce Boxeo y está construida para escalar a cualquier "
    "arte marcial o deporte de contacto sin reescribir el sistema.")
body(doc,
    "El producto ya está en producción: cualquier persona puede entrar desde un "
    "navegador, sin instalar nada, e iniciar sesión con una cuenta de demostración. "
    "Funciona también como aplicación web instalable (PWA) en el móvil.")

h3(doc, "1.1 Qué es CombatIQ")
bullet(doc, "Una plataforma web multi-rol (atleta, coach, dirección) donde se registra el estado diario del deportista y se visualiza el rendimiento del equipo completo.")
bullet(doc, "Un sistema de check-in de bienestar que convierte respuestas subjetivas en un Wellness Score objetivo del 0 al 100, ponderado por deporte.")
bullet(doc, "Un motor de análisis de señales biomédicas: frecuencia cardíaca y variabilidad (HRV) desde ECG, y detección de golpes, patadas e impactos desde sensores inerciales (IMU).")
bullet(doc, "Un módulo de análisis biomecánico por video (visión por computadora) que mide ángulos articulares, velocidad de pateo, rango de movimiento y simetría, e incluso analiza un combate rojo-vs-azul cuadro a cuadro.")
bullet(doc, "Un asistente de inteligencia artificial que lee todos esos datos y escribe análisis de coaching en lenguaje natural, adaptados al rol y al deporte.")
bullet(doc, "Un panel de coach que muestra el estado de todo el equipo de un vistazo, con alertas automáticas por niveles, historial y chat interno.")

h3(doc, "1.2 Qué hace CombatIQ hoy (estado real)")
bullet(doc, "En producción en la nube (Railway), con base de datos PostgreSQL gestionada (Supabase) y respaldo dual con SQLite para desarrollo.")
bullet(doc, "Login funcional con cuentas demo (atleta TKD, coach TKD, dirección) para probar el producto de inmediato.")
bullet(doc, "Check-in de bienestar, dashboard de coach, análisis ECG/IMU, replay de combate, biomecánica por video y exportación de informes profesionales (PDF, Excel, CSV) operativos.")
bullet(doc, "Análisis biomecánico validado contra valores de referencia de élite (World Taekwondo) sobre video real de competencia.")

h3(doc, "1.3 Qué NO es (todavía) — límites honestos")
bullet(doc, "No es todavía una app móvil nativa de App Store / Google Play; es web instalable (PWA) optimizada para móvil.")
bullet(doc, "No transmite aún los sensores en tiempo real por Bluetooth de forma comercial: la ingestión de ECG/IMU es por archivo (CSV/JSON) o por API; el hub BLE existe como prototipo funcional validado en modo demo.")
bullet(doc, "El análisis biomecánico por video se hace offline sobre grabaciones, no en vivo durante el combate (es computacionalmente intenso).")
bullet(doc, "No reemplaza al coach ni emite diagnóstico médico: es una herramienta de apoyo a la decisión; el criterio profesional sigue mandando.")
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 2. ALCANCE FUNCIONAL DETALLADO
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "2. Alcance funcional detallado")

h3(doc, "2.1 Flujo del deportista")
numbered(doc, "Inicia sesión en la plataforma desde el navegador del móvil o la computadora (o la instala como app PWA).")
numbered(doc, "La primera vez completa un onboarding breve: deporte, datos básicos y foto de perfil.")
numbered(doc, "Cada mañana, antes de entrenar, hace su check-in de bienestar: responde en menos de dos minutos sliders sobre sueño, energía, fatiga, molestias y disposición, más el esfuerzo (RPE) y la duración de la sesión anterior.")
numbered(doc, "El sistema calcula su Wellness Score (0-100) y le muestra su tendencia de los últimos días con zonas de color.")
numbered(doc, "Si entrenó con sensores, sube el archivo de su banda de pecho (ECG) o de sus IMU y ve sus métricas: frecuencia cardíaca, variabilidad, golpes, patadas, intensidad.")
numbered(doc, "Si grabó su combate o entrenamiento, sube el video y obtiene su análisis biomecánico: ángulos, velocidad de pateo, simetría y una lectura táctica.")
numbered(doc, "Consulta el historial de sus sesiones, su peso y su control de categoría de competencia.")
numbered(doc, "Recibe del coach (y de la IA) recomendaciones concretas sobre qué trabajar y se comunica por el chat interno.")

h3(doc, "2.2 Flujo del coach")
numbered(doc, "Inicia sesión y ve el dashboard de su equipo: todos sus atletas del deporte que dirige, cada uno con su semáforo de estado (verde / ámbar / rojo).")
numbered(doc, "Detecta de un vistazo quién llega listo y quién necesita atención, sin abrir conversaciones una por una.")
numbered(doc, "Entra al perfil de un atleta y revisa su tendencia de bienestar, su carga (ACWR), su recuperación (HRV) y sus sesiones recientes.")
numbered(doc, "Revisa las señales del atleta: ECG/HRV, volumen e intensidad de impactos, replay de combate sincronizado con video.")
numbered(doc, "Lee la nota de coaching que genera la IA para cada atleta y sus alertas automáticas (fatiga, sobrecarga, bienestar bajo).")
numbered(doc, "Analiza un combate en modo rojo-vs-azul: la plataforma sigue a ambos atletas, mide su técnica y entrega una lectura táctica comparada.")
numbered(doc, "Crea sesiones, asigna sensores a sus atletas, exporta informes individuales o de equipo en PDF/Excel y se comunica por chat.")
numbered(doc, "Solo ve atletas de su propio deporte (un coach de TKD ve TKD; uno de Boxeo ve Boxeo); los datos están separados por disciplina.")

h3(doc, "2.3 Flujo de la dirección / administración")
numbered(doc, "Inicia sesión con rol de dirección (admin) o de observador (inversor).")
numbered(doc, "Visualiza métricas globales del programa: atletas activos, sesiones, distribución por deporte y estado general.")
numbered(doc, "Da de alta, edita o desactiva usuarios y revisa la asignación de atletas a cada coach.")
numbered(doc, "Consulta el panel ejecutivo con la visión agregada del rendimiento y la actividad.")
numbered(doc, "Exporta reportes para seguimiento y rendición de cuentas.")
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 3. STACK TECNOLOGICO
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "3. Stack tecnológico")
body(doc,
    "El stack actual prioriza tres cosas: (1) tener un producto real y demostrable "
    "en producción, (2) que cada pieza de análisis sea explicable y defendible, y "
    "(3) un camino de escalado claro sin reescribir lo que ya funciona.")
add_table(doc,
    ["Capa", "Tecnología", "Por qué"],
    [
        ["Frontend",
         "Plotly Dash 2.16 + Plotly 6.4 + CSS propio + JavaScript",
         "Componentes reactivos y gráficas biomédicas interactivas en un solo lenguaje (Python). Tema oscuro/claro y PWA instalable."],
        ["Backend",
         "Flask 3.0 + Python 3.11 (servidor Dash embebido)",
         "Autenticación (bcrypt), sesiones, navegación protegida por rol y API REST interna para ingestión de sensores."],
        ["Base de datos",
         "PostgreSQL gestionado (Supabase) en producción / SQLite (WAL) en local",
         "Misma base de código para ambos motores: migraciones versionadas e idempotentes (v10 → v210), 25 tablas."],
        ["Visión por computadora",
         "MediaPipe (pose) + YOLOv8-pose + OpenVINO + ByteTrack + OpenCV",
         "Estimación de pose y seguimiento multi-atleta para el análisis biomecánico por video."],
        ["Inteligencia artificial",
         "API de Claude (Anthropic) — modelos Haiku y Opus",
         "Genera análisis de coaching en lenguaje natural, adaptado al rol y al deporte, con degradación elegante si no hay API key."],
        ["Procesamiento de señal",
         "NumPy + SciPy (scipy.signal)",
         "Detección de picos R en ECG, cálculo de HRV (SDNN, RMSSD) y detección de impactos en señales IMU."],
        ["Informes / exports",
         "ReportLab (PDF) + openpyxl (Excel) + Kaleido (gráficas)",
         "Informes profesionales de atleta y equipo en PDF, hojas Excel y CSV con unidades y metadatos."],
        ["Hosting / Deploy",
         "Railway (auto-deploy desde GitHub) + gunicorn",
         "Despliegue continuo, contenedor Linux reproducible (Dockerfile), endpoint de salud y reinicio automático."],
        ["Control de versiones",
         "GitHub (repositorio privado)",
         "Historial completo, respaldo del código e integración directa con el despliegue."],
        ["Hardware / sensores",
         "Hub BLE en Python (Bleak) + ESP32 + IMU MPU-6050",
         "Puente entre los sensores físicos y la plataforma vía Bluetooth (Nordic UART) y API REST."],
    ],
    col_widths=[3.3, 5.2, 7.5]
)
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 4. MODELO DE DATOS
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "4. Modelo de datos (alto nivel)")
body(doc,
    "La base de datos organiza el sistema alrededor de una entidad central —la "
    "sesión de entrenamiento— que enlaza el check-in, las señales y el video de un "
    "mismo día. Estas son las tablas principales (el esquema completo tiene 25 tablas).")
add_table(doc,
    ["Tabla", "Para qué sirve"],
    [
        ["users", "Todos los usuarios: atleta, coach, dirección. Guarda rol, deporte, foto y a qué coach pertenece."],
        ["questionnaires", "Cada check-in diario: respuestas, Wellness Score (0-100), RPE y duración de la sesión anterior."],
        ["sessions", "La sesión de entrenamiento. Agrupa ECG, IMU y cuestionario bajo una misma fecha y deporte."],
        ["ecg_files / ecg_metrics", "Archivos de señal cardíaca y sus métricas calculadas: BPM, SDNN, RMSSD, latidos detectados."],
        ["imu_metrics", "Métricas de impacto por sensor: golpes/patadas, cadencia, intensidad (g) y velocidad angular."],
        ["sensor_sessions", "Actividad de sensores ligada a una sesión (preparada para streaming en tiempo real)."],
        ["coach_athletes / teams", "Relación coach-atleta y gestión de equipos por deporte."],
        ["weights", "Historial de peso y peso objetivo para control de categoría de competencia."],
        ["messages", "Chat interno coach-atleta, con marca de leído/no leído."],
        ["audit / migrations", "Trazabilidad de cambios de esquema versionados e idempotentes."],
    ],
    col_widths=[4.5, 11.5]
)

h3(doc, "4.1 Roles y permisos")
bullet(doc, "Atleta: ve y gestiona solo sus propios datos, sesiones, sensores y análisis.")
bullet(doc, "Coach: ve únicamente a los atletas de su deporte; no accede a atletas de otras disciplinas ni de otros coaches.")
bullet(doc, "Dirección (admin): visión global, alta/baja de usuarios y asignaciones.")
bullet(doc, "Observador (inversor): acceso de solo lectura a la visión ejecutiva, útil para demostraciones.")
bullet(doc, "Toda exportación de informes valida la sesión y los permisos en el servidor antes de generar el archivo.")
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 5. EL ANALISIS QUE HAY DETRAS
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "5. El análisis que hay detrás")
body(doc,
    "Esta es la parte que diferencia a CombatIQ de un simple formulario o un "
    "visualizador de sensores. Cada dato se procesa con métodos transparentes y se "
    "interpreta en contexto de los demás. A continuación, cómo funciona cada motor.")

h3(doc, "5.1 Wellness Score — el check-in inteligente")
body(doc,
    "El cuestionario tiene 22 preguntas organizadas en 7 grupos (base, taekwondo, "
    "boxeo, competencia, peso y molestias). Cada pregunta tiene un peso (3-14) y una "
    "dimensión: «positiva» (más es mejor, como el sueño) o «de riesgo» (más es peor, "
    "como el dolor). El motor normaliza cada respuesta a una escala 0-100, invierte "
    "las de riesgo y pondera cada una según su importancia para el deporte del atleta. "
    "El resultado es un único número, el Wellness Score, que resume qué tan listo "
    "está el atleta para entrenar hoy. No es un promedio simple: el sueño pesa más "
    "que una molestia leve, y eso está calibrado por disciplina.")

h3(doc, "5.2 ECG / HRV — el estado del corazón")
body(doc,
    "Desde la señal de una banda de pecho, el sistema detecta los picos R del "
    "electrocardiograma (con scipy.signal.find_peaks), mide los intervalos entre "
    "latidos (R-R) y de ahí calcula:")
add_table(doc,
    ["Métrica", "Qué mide", "Cómo se calcula"],
    [
        ["BPM", "Frecuencia cardíaca", "60 / promedio de los intervalos R-R en segundos."],
        ["SDNN", "Variabilidad global (HRV)", "Desviación estándar de todos los intervalos R-R. Valores bajos = fatiga acumulada."],
        ["RMSSD", "Recuperación autonómica", "Raíz del promedio de las diferencias al cuadrado entre latidos consecutivos."],
        ["n_peaks", "Calidad de la señal", "Número de latidos detectados; valida que la grabación sea fiable."],
    ],
    col_widths=[2.6, 4.6, 8.8]
)
body(doc,
    "La variabilidad cardíaca (HRV) es, hoy, uno de los indicadores más precisos de "
    "la ciencia del deporte para saber si un atleta está recuperado o sobreentrenado.")

h3(doc, "5.3 IMU — golpes, patadas e impactos")
body(doc,
    "Los sensores inerciales (acelerómetro + giroscopio) miden el movimiento. "
    "El sistema calcula la magnitud del impacto (raíz de ax²+ay²+az²) y, según la "
    "posición del sensor, aplica un umbral distinto para contar el evento: golpe de "
    "muñeca, patada de pie o impacto recibido en la cabeza.")
add_table(doc,
    ["Métrica", "Qué mide"],
    [
        ["Número de impactos", "Total de golpes, patadas o impactos detectados sobre el umbral."],
        ["Cadencia (por minuto)", "Volumen de trabajo: impactos divididos por la duración."],
        ["Intensidad (g)", "Fuerza media y máxima de cada impacto, en gravedades."],
        ["Velocidad angular", "Rapidez de rotación del segmento corporal, derivada del giroscopio."],
        ["Simetría", "Ratio entre lado dominante y no dominante (clave en taekwondo)."],
    ],
    col_widths=[4.0, 12.0]
)

h3(doc, "5.4 Carga interna/externa y ACWR")
body(doc,
    "El motor de análisis agrega el historial del atleta y calcula la carga de "
    "trabajo. El ACWR (ratio de carga aguda vs crónica) compara las últimas 7 "
    "sesiones contra las últimas 28: si la carga reciente sube demasiado respecto a "
    "lo habitual, hay riesgo de sobrecarga. La readiness por HRV compara la "
    "recuperación del día contra la línea base de 30 días del propio atleta. De ahí "
    "salen las alertas automáticas clasificadas en tres niveles: ok, atención "
    "(warning) y riesgo (danger).")

h3(doc, "5.5 Biomecánica por video — visión por computadora")
body(doc,
    "Es el pilar diferenciador de CombatIQ. A partir de un video del entrenamiento o "
    "combate, sin ningún sensor adicional, el sistema reconstruye el movimiento:")
bullet(doc, "Estimación de pose con MediaPipe y YOLOv8-pose: detecta las articulaciones de cada atleta cuadro a cuadro.")
bullet(doc, "Seguimiento multi-atleta con ByteTrack: mantiene la identidad de cada peleador y clasifica peto rojo vs peto azul por color.")
bullet(doc, "Métricas de técnica: ángulos articulares, rango de movimiento (ROM) de rodilla y cadera, simetría bilateral y velocidad angular.")
bullet(doc, "Indicadores específicos de taekwondo: velocidad de pateo estimada en m/s y ángulo de cámara de la pierna de golpeo (chamber angle).")
bullet(doc, "Modo combate rojo-vs-azul: analiza a ambos atletas en simultáneo, mide la distancia entre ellos y entrega una lectura táctica comparada.")
body(doc,
    "Los resultados se validaron contra valores de referencia de atletas de élite de "
    "World Taekwondo (por ejemplo, velocidad de pateo de 10-17 m/s) usando video real "
    "de competencia. Regla de oro del sistema: si una medición supera lo que un humano "
    "puede hacer físicamente, se trata como error de calibración, no como un atleta "
    "excepcional. El procesamiento corre en modo offline (aproximadamente 3x el tiempo "
    "real del video en un equipo sin GPU dedicada).")

h3(doc, "5.6 Inteligencia artificial contextual")
body(doc,
    "Toda esa información —wellness, carga, recuperación, impactos, biomecánica— la "
    "lee un modelo de IA (Claude, de Anthropic). En lugar de dejar tablas de números, "
    "escribe un párrafo en lenguaje natural explicando el estado del atleta y qué "
    "cuidados tomar. El análisis es consciente del rol (le habla distinto al atleta y "
    "al coach) y del deporte. Incluye caché, timeout y degradación elegante: si el "
    "servicio de IA no está disponible, el resto de la plataforma sigue funcionando "
    "con normalidad. La IA no sustituye al coach; es un asistente que no se cansa, no "
    "olvida nada y siempre tiene los datos frescos.")
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 6. HARDWARE Y SENSORES
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "6. Hardware y sensores")
body(doc,
    "El hardware es el diferenciador a largo plazo del producto. La plataforma ya "
    "tiene el catálogo de sensores, la asignación por atleta y la infraestructura de "
    "ingestión; el siguiente paso es el despliegue del hardware físico a escala.")

h3(doc, "6.1 Catálogo de sensores")
add_table(doc,
    ["Código", "Sensor", "Deporte", "Métricas clave", "Estado"],
    [
        ["ECG", "Banda torácica ECG/HRV", "TKD · Boxeo", "BPM, SDNN, RMSSD, latidos", "CSV/API listo"],
        ["IMU_WRIST", "IMU muñeca / guante", "TKD · Boxeo", "golpes, golpes/min, intensidad (g)", "BLE/API/CSV listo"],
        ["IMU_FOOT", "IMU tobillo / pie", "TKD", "patadas, potencia (g), simetría", "BLE/API/CSV listo"],
        ["IMU_HEAD", "IMU casco / cabeza", "TKD · Boxeo", "impactos, pico g, impactos >3g/>6g", "BLE/API/CSV listo"],
        ["HR_WRIST", "FC reloj / pulsera", "TKD · Boxeo", "FC media/máxima, zonas", "Experimental (CSV/API)"],
    ],
    col_widths=[2.3, 3.4, 2.6, 4.6, 3.1]
)

h3(doc, "6.2 Arquitectura del hub BLE")
body(doc,
    "Un sensor de combate envía sus datos a un hub (un programa puente, escrito en "
    "Python) que los reenvía a la plataforma. El recorrido es:")
mono(doc,
"  Sensor fisico            Hub BLE (PC/movil)            CombatIQ\n"
"  -------------            ------------------            --------\n"
"  ESP32 + MPU-6050   --->  Bleak escucha el     --->     POST /api/sensor-data\n"
"  (acel. + giro)           servicio Nordic UART          (cada ~10 s)\n"
"                           detecta picos de impacto,\n"
"                           clasifica golpe/patada/\n"
"                           impacto y mide vel. angular")
bullet(doc, "Microcontrolador ESP32 con IMU MPU-6050 (acelerómetro a ±2g, giroscopio a ±250°/s) anunciado como dispositivo «CombatIQ-*».")
bullet(doc, "Comunicación por Bluetooth Low Energy usando el Nordic UART Service (NUS).")
bullet(doc, "Umbrales de impacto por posición: 2.5 g en muñeca/guante, 3.5 g en pie, 1.5 g en casco (más sensible para impactos recibidos).")
bullet(doc, "El hub corre en modo demo (sin hardware) para validar el pipeline completo de extremo a extremo, y envía a la API REST de la plataforma.")

h3(doc, "6.3 Estado del pipeline físico — 3 fases")
add_table(doc,
    ["Fase", "Alcance", "Estado"],
    [
        ["Catálogo + ingestión", "Catálogo de sensores, asignación por atleta, API REST y subida CSV/JSON", "Listo"],
        ["Fase 1 — IMU en vivo", "Streaming BLE de golpes/patadas en tiempo real, ligado a la sesión activa", "Pipeline validado en demo; pendiente hardware a escala"],
        ["Fase 2 — ECG en vivo", "Streaming de ECG y HRV en vivo (hoy por archivo/API)", "Planificado"],
        ["Fase 3 — EMG", "Activación y fatiga muscular por round", "Futuro"],
    ],
    col_widths=[3.6, 8.4, 4.0]
)
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 7. ESTADO ACTUAL DEL PRODUCTO
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "7. Estado actual del producto")
body(doc,
    "Resumen honesto de qué está operativo hoy, qué está en progreso y qué queda "
    "deliberadamente fuera del alcance actual.")

h3(doc, "7.1 Operativo en producción")
bullet(doc, "App desplegada en la nube con base de datos PostgreSQL gestionada y respaldo dual SQLite.")
bullet(doc, "Autenticación segura por rol (bcrypt), navegación protegida y cuentas demo funcionales.")
bullet(doc, "Check-in de bienestar con Wellness Score ponderado por deporte y tendencia con zonas de color.")
bullet(doc, "Dashboard de coach con semáforo de equipo, alertas automáticas y filtro por deporte.")
bullet(doc, "Análisis ECG/HRV e IMU, replay de combate con video sincronizado y eventos detectados.")
bullet(doc, "Biomecánica por video (pose, ROM, simetría, velocidad de pateo, modo rojo-vs-azul) validada con video real.")
bullet(doc, "Notas de coaching con IA, conscientes de rol y deporte, con degradación elegante.")
bullet(doc, "Exportación de informes profesionales en PDF, Excel y CSV; chat interno; control de peso; gestión de equipos.")
bullet(doc, "Aplicación instalable como PWA, con modo claro y oscuro.")

h3(doc, "7.2 En progreso / próximo")
bullet(doc, "Pulido pixel-perfect de todas las pantallas de Taekwondo de cara a la demostración de julio de 2026.")
bullet(doc, "Sistema de diseño en Figma con paridad código-diseño (pantallas de coach y biomecánica).")
bullet(doc, "Responsive móvil-first real para el coach en el gimnasio.")
bullet(doc, "Dominio propio y optimización de la PWA.")
bullet(doc, "Hardware IMU real conectado al hub a escala (Fase 1 del pipeline de sensores).")

h3(doc, "7.3 Fuera del alcance actual (deliberado)")
bullet(doc, "Apps nativas iOS/Android en tiendas (la PWA cubre el caso de uso hoy).")
bullet(doc, "Streaming biomecánico en vivo durante el combate (el análisis es offline).")
bullet(doc, "Expansión multi-deporte completa (Karate, Judo, Kickboxing, MMA) — se escala después de validar Taekwondo.")
bullet(doc, "EMG y sensores de respiración en vivo — fases posteriores del roadmap de hardware.")
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 8. RIESGOS Y MITIGACIONES
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "8. Riesgos y mitigaciones")
add_table(doc,
    ["Riesgo", "Impacto", "Mitigación"],
    [
        ["Disponibilidad de hardware de sensores a escala",
         "Alto",
         "El software ya funciona con datos por archivo/API y demo; el hardware se incorpora por fases sin bloquear el producto."],
        ["Costo de cómputo del análisis por video (sin GPU)",
         "Medio",
         "Procesamiento offline optimizado (~3x tiempo real) y separado en pasos; ruta a GPU/servidor dedicado cuando el volumen lo justifique."],
        ["Dependencia de un servicio externo de IA",
         "Medio",
         "Caché, timeout y degradación elegante: la app funciona completa aunque la IA no responda."],
        ["Concurrencia de la base de datos al crecer",
         "Medio",
         "Ya en PostgreSQL gestionado; misma base de código soporta el escalado sin reescritura."],
        ["Producto mantenido por un equipo pequeño",
         "Alto",
         "Código versionado en GitHub, despliegue continuo, migraciones idempotentes y documentación técnica continua."],
        ["Que un competidor copie la idea",
         "Medio",
         "La ventaja no es solo el código: es el conocimiento específico del deporte, los datos acumulados y la relación con coaches."],
    ],
    header_hex="B02A2A",
    col_widths=[5.0, 2.2, 8.8]
)
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 9. ESCALABILIDAD Y PLANES FUTUROS
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "9. Escalabilidad y planes futuros")

h3(doc, "9.1 Escalabilidad técnica")
body(doc,
    "La plataforma está lista para crecer sin rehacerse. La base de datos ya es "
    "PostgreSQL gestionado y la misma base de código soporta SQLite en local. El "
    "camino de evolución (V2), una vez validado el mercado, separa el frontend del "
    "backend: una interfaz moderna (React + Tailwind) sobre una API (FastAPI), "
    "manteniendo intactos los motores que ya funcionan (análisis de señal, "
    "biomecánica e IA). Mejoras previstas: caché compartida (Redis), actualización "
    "en tiempo real del dashboard (WebSockets) y notificaciones push al móvil del "
    "coach ante un wellness bajo o un impacto de alto riesgo.")

h3(doc, "9.2 Escalabilidad de producto — multi-deporte")
body(doc,
    "CombatIQ está construida por deporte: las preguntas del check-in, las métricas "
    "de sensores, los umbrales de alerta y el lenguaje de la IA son configurables por "
    "disciplina. Hoy el foco es Taekwondo (donde se validó con datos reales) y el "
    "sistema ya reconoce Boxeo. Agregar Karate, Judo, Kickboxing o MMA no es "
    "reescribir la plataforma: es calibrar parámetros. Esa es la diferencia entre un "
    "producto escalable y uno que hay que reconstruir cada vez.")

h3(doc, "9.3 Roadmap de hardware")
body(doc,
    "El plan de sensores avanza en fases: primero IMU en vivo (golpes y patadas en "
    "tiempo real ligados a la sesión), luego ECG/HRV en streaming, y finalmente EMG "
    "para activación y fatiga muscular. La diferenciación por deporte también aplica "
    "al hardware: guante IMU para boxeo, tobillera para taekwondo, casco para impactos "
    "recibidos. El software se completa primero; el hardware se incorpora encima sin "
    "romper lo que ya funciona.")

h3(doc, "9.4 Por qué es difícil de copiar")
numbered(doc, "Integra tres mundos —bienestar subjetivo, sensores biomédicos y biomecánica por video— en un sistema coherente donde cada dato se entiende en contexto.")
numbered(doc, "Está calibrada específicamente para deportes de combate; un coach siente que la herramienta entiende su deporte, no que es un fitness tracker genérico.")
numbered(doc, "Usa IA para generar coaching en lenguaje natural, algo que no existe en herramientas específicas de combate en el mercado hispanohablante.")
numbered(doc, "Los datos son del equipo: historial propio de rendimiento, no información atrapada en servidores ajenos.")
numbered(doc, "Previene en lugar de reaccionar: detecta patrones de sobrecarga antes de la lesión, no después.")
add_hr(doc)
section_break(doc)

# ══════════════════════════════════════════════════════════════════════════
# 10. RESUMEN EN 5 MINUTOS (EN PALABRAS SIMPLES)
# ══════════════════════════════════════════════════════════════════════════
h2(doc, "10. Resumen en 5 minutos (en palabras simples)")
body(doc,
    "Un coach de Taekwondo o Boxeo hoy trabaja a ciegas: pregunta «¿cómo te "
    "sientes?» por WhatsApp, anota en papel, sube datos de un sensor a Excel y trata "
    "de interpretarlos sin contexto. Si dirige 15 atletas, esa información está "
    "dispersa en 15 conversaciones. CombatIQ reemplaza todo eso con una sola "
    "plataforma, ordenada y automática.")
body(doc,
    "Antes de entrenar, el atleta responde unas preguntas en su celular. La app las "
    "convierte en un número del 0 al 100 que el coach ve al instante. Si el atleta "
    "usó una banda de pecho, la app calcula su frecuencia cardíaca y su variabilidad. "
    "Si llevó un sensor en la muñeca o el tobillo, cuenta cuántos golpes dio, a qué "
    "ritmo y con qué fuerza. Si grabó su combate, mide su técnica desde el video. "
    "Y al final, la inteligencia artificial lee todo junto y le escribe al coach, "
    "en lenguaje natural, cómo está el atleta y qué cuidar en el próximo entrenamiento.")

add_table(doc,
    ["Lo que se usa hoy", "Su problema", "Cómo lo resuelve CombatIQ"],
    [
        ["Papel / WhatsApp", "Sin historial ni análisis; la información se pierde.",
         "Registro digital automático, con historial por sesión y por atleta."],
        ["Polar / Garmin solos", "Solo miden ritmo cardíaco, sin contexto.",
         "Añade bienestar, impactos, biomecánica e IA encima del dato del corazón."],
        ["TrainingPeaks / genéricos", "No son para combate; sin detección de golpes.",
         "Cuestionarios, alertas y métricas calibradas para TKD y Boxeo."],
        ["Excel del profesor", "Manual, lento, sin alertas.",
         "Los datos entran solos y las alertas salen solas."],
    ],
    col_widths=[3.8, 5.0, 7.2]
)

body(doc,
    "En una frase: CombatIQ convierte la intuición del coach en decisiones con datos, "
    "para que el atleta llegue al campeonato en su mejor forma —y no lesionado—.")

add_hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Documento de visión de producto · Mayo 2026 · CombatIQ — Monitoreo para Deportes de Combate")
run.font.size = Pt(8.5)
run.font.color.rgb = GRAY


# ── Guardar ──────────────────────────────────────────────────────────────────
import os, shutil
out_tmp   = "Alcance_CombatIQ_new.docx"
out_final = "Alcance_CombatIQ.docx"
doc.save(out_tmp)
if os.path.exists(out_final):
    try:
        os.remove(out_final)
    except PermissionError:
        print("AVISO: cierra el Word antes de reemplazar el archivo original.")
        print("Archivo guardado como: " + out_tmp)
        raise SystemExit(0)
shutil.move(out_tmp, out_final)
print("Archivo generado: " + out_final)
